"""
Script to generate the Word (.docx) editable version of
Fundamentacion_Metodologica_Digital_Twins_inDrive.docx
with CRIBA Research styling.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_docx(filename="Fundamentacion_Metodologica_Digital_Twins_inDrive.docx"):
    doc = Document()
    
    # Page Setup
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    # Header / Footer setup
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.text = "CRIBA RESEARCH • Digital Customer Twins (inDrive Perú)"
    header_p.runs[0].font.name = 'Calibri'
    header_p.runs[0].font.size = Pt(8.5)
    header_p.runs[0].font.bold = True
    header_p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    footer_p = section.footer.paragraphs[0]
    footer_p.text = "Confidencial • Preparado para CRIBA Research & inDrive"
    footer_p.runs[0].font.name = 'Calibri'
    footer_p.runs[0].font.size = Pt(8.5)
    footer_p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Document Header Title
    p_title = doc.add_paragraph()
    r_title = p_title.add_run("FUNDAMENTACIÓN METODOLÓGICA Y TRAZABILIDAD")
    r_title.font.name = 'Arial'
    r_title.font.size = Pt(18)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    p_title.paragraph_format.space_after = Pt(2)

    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("Digital Customer Twins • Estudio Cualitativo de Mototaxis (inDrive, Lima)")
    r_sub.font.name = 'Arial'
    r_sub.font.size = Pt(11)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)
    p_sub.paragraph_format.space_after = Pt(10)

    # Metadata Table
    meta_table = doc.add_table(rows=3, cols=4)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    meta_data = [
        [("Proyecto:", True), ("Digital Customer Twins (inDrive Perú)", False), ("Fecha:", True), ("Septiembre 2026", False)],
        [("Metodología:", True), ("Grounded Qualitative AI & RAG Local", False), ("Muestra:", True), ("25 Entrevistas en Profundidad", False)],
        [("Objetivo:", True), ("Sustento analítico y trazabilidad empírica", False), ("Estado:", True), ("Entregable Técnico de Validación", False)]
    ]

    col_widths = [Inches(1.1), Inches(2.7), Inches(0.9), Inches(2.3)]

    for row_idx, row in enumerate(meta_table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.width = col_widths[col_idx]
            label, is_bold = meta_data[row_idx][col_idx]
            set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(label)
            r.font.size = Pt(9)
            r.font.bold = is_bold
            r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) if is_bold else RGBColor(0x47, 0x55, 0x69)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Helper for Section Titles
    def add_section_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        return p

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        r_b = p.add_run(bold_prefix + " ")
        r_b.bold = True
        r_b.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        r_t = p.add_run(text)
        return p

    # Section 1
    add_section_h2("1. Marco Metodológico y Procesamiento del Corpus")
    p1 = doc.add_paragraph("El desarrollo del sistema de Digital Customer Twins responde a la necesidad de transformar el material cualitativo de campo en un modelo analítico interactivo y explorable en tiempo real, superando el carácter estático de las presentaciones tradicionales. El sistema no genera conocimiento abstracto ni especulativo; opera strictly sobre la evidencia recolectada en el trabajo de campo.")
    p1.paragraph_format.space_after = Pt(6)

    add_bullet("Base empírica primaria:", "El corpus del sistema está constituido por 25 transcripciones integrales de entrevistas en profundidad realizadas a conductores de mototaxis en zonas clave de Lima (Comas, Collique, Añashuayco), además de las guías de moderación y el informe de investigación.")
    add_bullet("Estructuración e ingesta:", "Cada entrevista fue fragmentada e indexada preservando metadatos clave como ID de participante, distrito, antigüedad en el rubro y temas abordados.")
    add_bullet("Matriz de evidencia estructurada:", "Se construyó una matriz relacional que vincula temas críticos (tarifas, comisiones, zonas peligrosas, bonos) con los testimonios empíricos de los conductores.")

    # Section 2
    add_section_h2("2. Calibración de los Digital Customer Twins")
    p2 = doc.add_paragraph("A partir del análisis inductivo de las transcripciones, se sintetizaron tres perfiles sintéticos diferenciados que representan los arquetipos dominantes de la muestra:")
    p2.paragraph_format.space_after = Pt(6)

    # Twins Table
    twins_table = doc.add_table(rows=4, cols=3)
    twins_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_widths = [Inches(1.8), Inches(2.2), Inches(3.0)]

    headers = ["Perfil Sintético", "Foco Operativo y Actitudinal", "Patrón de Comportamiento Dominante"]
    for idx, cell in enumerate(twins_table.rows[0].cells):
        cell.width = t_widths[idx]
        set_cell_background(cell, "F0FDF4")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(headers[idx])
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x06, 0x5F, 0x46)

    row_data = [
        ("Twin A\nDisciplined Hard Work\n(Marcos)", "Control de margen neto diario y predictibilidad.", "Exige conocer el destino antes de aceptar la carrera. Rechaza viajes a zonas peligrosas (cerros/sin luz) y prioriza tarifas justas sobre volumen."),
        ("Twin B\nTactical Cash Optimizer\n(Julio)", "Velocidad de rotación y flujo continuo de caja.", "Orientado a la acumulación rápida de efectivo. Maximiza bonos diarios por cuota de viajes y negocia tarifas de forma agresiva."),
        ("Twin C\nLow-Pressure Flexibles\n(Carlos)", "Flexibilidad horaria y reducción del estrés.", "Combina el paradero físico tradicional con solicitudes de la app. Trabaja a su propio ritmo y evita jornadas extensas o competitivas.")
    ]

    for r_idx, (c1_txt, c2_txt, c3_txt) in enumerate(row_data):
        row_cells = twins_table.rows[r_idx + 1].cells
        for col_idx, text_val in enumerate([c1_txt, c2_txt, c3_txt]):
            cell = row_cells[col_idx]
            cell.width = t_widths[col_idx]
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text_val)
            r.font.size = Pt(9)
            if col_idx == 0:
                r.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Section 3
    add_section_h2("3. Arquitectura de Anclaje a la Evidencia (Grounding y RAG Local)")
    p3 = doc.add_paragraph("El principal diferencial metodológico de la herramienta es su mecanismo de anclaje (grounding), diseñado para prevenir alucinaciones o respuestas genéricas fuera de contexto:")
    p3.paragraph_format.space_after = Pt(6)

    add_bullet("1. Recuperación Semántica Acotada:", "Ante cada consulta del usuario, un motor de búsqueda cualitativa (TF-IDF + Cosine Similarity) extrae de forma prioritaria los fragmentos de transcripción más pertinentes del corpus.")
    add_bullet("2. Trazabilidad Explícita:", "Cada respuesta emitida por un Twin incluye una ventana de trazabilidad que muestra la cita textual exacta, el participante y el código de la entrevista de origen (ej. EP01, EP05, EP18).")
    add_bullet("3. Detección de Límites de Información:", "Si se le formula una pregunta al Twin sobre un tema no abordado en el trabajo de campo (ej. vehículos eléctricos), el sistema emite una advertencia formal notificando que el material disponible no contiene evidencia suficiente.")

    # Section 4
    add_section_h2("4. Modos de Análisis y Aplicabilidad Estratégica")
    add_bullet("Modo Chat Individual:", "Permite profundizar en la psicología y motivaciones de un perfil específico frente a preguntas puntuales.")
    add_bullet("Modo Matriz Side-by-Side:", "Permite ingresar un cambio de propuesta comercial (ej. modificación de comisión o filtro de seguridad) y contrastar la reacción simultánea de los perfiles lado a lado con disparadores de 1-clic.")
    add_bullet("Modo Focus Group Interactivo:", "Orquesta un debate simulado entre los tres Twins sobre problemáticas complejas con síntesis cualitativa del moderador.")
    add_bullet("Modo Cuadros & Matriz:", "Ofrece tablas sintetizadas de 7 dimensiones estratégicas y comparativa de marcas exportables a Excel/CSV.")

    # Section 5 (NEW SECTION ALIGNED WITH USER REQUEST)
    add_section_h2("5. Protocolos de Confidencialidad, Gobernanza de Datos y Seguridad en IA")
    p5 = doc.add_paragraph("Garantizar el cumplimiento de las normativas de protección de datos y la salvaguarda de secretos comerciales constituye un pilar estructural de la plataforma. La arquitectura de seguridad del proyecto contempla dos niveles operativos actuales (MVP) y dos vías de escalamiento corporativo para inDrive:")
    p5.paragraph_format.space_after = Pt(6)

    add_bullet("1. Anonimización en Origen (Nivel MVP Aplicado):", "El 100% de los datos de identificación personal (PII) —nombres de entrevistados, números de teléfono, DNI y direcciones exactas— fueron purgados de las transcripciones y reemplazados por códigos anonimizados estandarizados (ej. [PARTICIPANTE_01]), eliminando cualquier riesgo de filtración de datos personales.")
    add_bullet("2. Procesamiento Híbrido y Motor Cualitativo Local (Nivel MVP Aplicado):", "La plataforma cuenta con un motor cualitativo local propio (TF-IDF / Búsqueda Semántica local) que consulta el corpus directamente en el servidor sin depender de enviar información sensible a servicios externos de IA.")
    add_bullet("3. Escalamiento a APIs Empresariales (Opción Enterprise Futura):", "Para un despliegue corporativo a escala, la plataforma permite conectar conectores comerciales (Google Vertex AI / OpenAI Enterprise) amparados por contratos de no-retención de datos (Zero Data Retention).")
    add_bullet("4. Servidor Privado Dedicado On-Premise (Opción Enterprise Futura):", "En caso de que inDrive exija aislamiento absoluto sobre sus datos estratégicos, la solución permite montarse en un servidor privado dedicado (VPS/VPC en AWS, GCP o Azure) ejecutando modelos de código abierto (Open Source) en un entorno 100% Air-Gapped sin conexión externa.")

    # Callout Box (Principio de Rigurosidad)
    quote_table = doc.add_table(rows=1, cols=1)
    quote_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    q_cell = quote_table.rows[0].cells[0]
    q_cell.width = Inches(7.0)
    set_cell_background(q_cell, "ECFDF5")
    set_cell_margins(q_cell, top=120, bottom=120, left=150, right=150)
    
    qp = q_cell.paragraphs[0]
    qp.paragraph_format.space_after = Pt(0)
    qr1 = qp.add_run("Principio de Rigurosidad Quali-AI & Privacidad:\n")
    qr1.font.bold = True
    qr1.font.size = Pt(9.5)
    qr1.font.color.rgb = RGBColor(0x06, 0x5F, 0x46)
    
    qr2 = qp.add_run('"Un Digital Customer Twin protege la privacidad de los participantes en origen y garantiza la confidencialidad estratégica del cliente, permitiendo que los equipos de Producto exploren evidencia real sin exponer datos sensibles."')
    qr2.font.italic = True
    qr2.font.size = Pt(9)
    qr2.font.color.rgb = RGBColor(0x04, 0x78, 0x57)

    doc.save(filename)
    print(f"Successfully generated {filename}")

if __name__ == "__main__":
    create_docx()
