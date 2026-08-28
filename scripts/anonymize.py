"""
Script de Anonimización de Entrevistas / Transcripciones (.docx)
Reemplaza nombres reales de los entrevistados por identificadores neutros
([PARTICIPANTE_01], [CONDUCTOR_02], etc.) en el texto y renombras los archivos.
"""

import os
import sys
import re
import docx

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from src.ingestion import ingest_all

# Mapeo de anonimización por Episodio / Nombre
PARTICIPANTS_MAP = {
    "EP01": {"names": ["Carlos Enrique", "Carlos"], "label": "[PARTICIPANTE_01]", "app": "inDrive"},
    "EP02": {"names": ["Gerson Yman", "Yerzo", "Gerson"], "label": "[PARTICIPANTE_02]", "app": "inDrive"},
    "EP03": {"names": ["Luis Aguilar", "Luis"], "label": "[PARTICIPANTE_03]", "app": "inDrive"},
    "EP04": {"names": ["Junior Cordova", "Junior"], "label": "[PARTICIPANTE_04]", "app": "inDrive"},
    "EP05": {"names": ["Jorla Vasquez", "Jorla"], "label": "[PARTICIPANTE_05]", "app": "Yango"},
    "EP06": {"names": ["Jorge Leonel", "Jorge"], "label": "[PARTICIPANTE_06]", "app": "Yango"},
    "EP07": {"names": ["John Anderson", "John"], "label": "[PARTICIPANTE_07]", "app": "General"},
    "EP08": {"names": ["Nicol Oros", "Nicol"], "label": "[PARTICIPANTE_08]", "app": "Yango"},
    "EP09": {"names": ["Maria Condori", "Maria"], "label": "[PARTICIPANTE_09]", "app": "Yango"},
    "EP10": {"names": ["Christopher Calero", "Christopher"], "label": "[PARTICIPANTE_10]", "app": "Uber"},
    "EP11": {"names": ["Jairton Reyes", "Jairton"], "label": "[PARTICIPANTE_11]", "app": "Yango"},
    "EP12": {"names": ["Daniel Rosales", "Daniel"], "label": "[PARTICIPANTE_12]", "app": "Yango"},
    "EP13": {"names": ["Jose Antonio Vaca", "Jose Antonio"], "label": "[PARTICIPANTE_13]", "app": "Yango"},
    "EP14": {"names": ["Juan Paul Ramirez", "Juan Paul"], "label": "[PARTICIPANTE_14]", "app": "Yango"},
    "EP15": {"names": ["Franco Ortiz", "Franco"], "label": "[PARTICIPANTE_15]", "app": "inDrive-Yango"},
    "EP16": {"names": ["Junior Sanchez", "Junior"], "label": "[PARTICIPANTE_16]", "app": "inDrive"},
    "EP17": {"names": ["Miguel Angel", "Miguel"], "label": "[PARTICIPANTE_17]", "app": "inDrive"},
    "EP18": {"names": ["Paul Vargas Ávalos", "Paul Vargas", "Paul"], "label": "[PARTICIPANTE_18]", "app": "inDrive"},
    "EP19": {"names": ["José Loaiza Minaya", "José Loaiza"], "label": "[PARTICIPANTE_19]", "app": "Yango"},
    "EP20": {"names": ["Jackelyn Bazalar Lara", "Jackelyn Bazalar", "Jackelyn"], "label": "[PARTICIPANTE_20]", "app": "inDrive"},
    "EP21": {"names": ["Jahiro Mitchell Martínez Maravi", "Jahiro Mitchell", "Jahiro"], "label": "[PARTICIPANTE_21]", "app": "Yango"},
    "EP22": {"names": ["John Díaz Zanabria", "John Díaz"], "label": "[PARTICIPANTE_22]", "app": "Yango"},
    "EP23": {"names": ["José calderon herrera", "José calderon"], "label": "[PARTICIPANTE_23]", "app": "inDrive"},
    "EP24": {"names": ["Azucena Engrario Carhuachin", "Azucena Engrario", "Azucena"], "label": "[PARTICIPANTE_24]", "app": "inDrive"},
    "EP25": {"names": ["José Eduardo Huynalaya Breña", "José Eduardo"], "label": "[PARTICIPANTE_25]", "app": "inDrive"}
}

def extract_ep_id(filename: str) -> str:
    match = re.search(r'(Ep|EP)\s*(\d+)', filename, re.IGNORECASE)
    if match:
        return f"EP{int(match.group(2)):02d}"
    return "EP00"

def anonymize_docx_file(filepath: str, ep_id: str, output_path: str):
    """Abre un .docx, anonimiza su texto y guarda el resultado."""
    import io
    with open(filepath, "rb") as f:
        file_stream = io.BytesIO(f.read())
    
    doc = docx.Document(file_stream)
    info = PARTICIPANTS_MAP.get(ep_id, {})
    names = info.get("names", [])
    label = info.get("label", f"[{ep_id}_PARTICIPANTE]")

    # Procesar párrafos
    for p in doc.paragraphs:
        if not p.text:
            continue
        text = p.text
        # Reemplazar encabezados que contengan el nombre del archivo / episodio
        for name in names:
            pattern = re.compile(re.escape(name), re.IGNORECASE)
            text = pattern.sub(label, text)
        p.text = text

    # Procesar tablas si existen
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text
                    for name in names:
                        pattern = re.compile(re.escape(name), re.IGNORECASE)
                        text = pattern.sub(label, text)
                    p.text = text

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    
    import time
    for attempt in range(5):
        try:
            with open(output_path, "wb") as f:
                f.write(buf.read())
            break
        except PermissionError:
            if attempt == 4:
                print(f"  [WARN] No se pudo sobrescribir {output_path} por bloqueo de sincronización (TeraBox/Cloud Drive), se omitió.")
            else:
                time.sleep(0.3)

def process_directory(target_dir: str):
    """Anonimiza y renombra todos los archivos .docx de un directorio."""
    if not os.path.exists(target_dir):
        return

    # Limpiar cualquier residuo de .tmp
    for f in os.listdir(target_dir):
        if f.endswith(".tmp"):
            try:
                os.remove(os.path.join(target_dir, f))
            except Exception:
                pass

    files = [f for f in os.listdir(target_dir) if f.endswith(".docx") and not f.startswith("~$")]
    
    for filename in files:
        filepath = os.path.join(target_dir, filename)
        ep_id = extract_ep_id(filename)
        if ep_id == "EP00":
            continue

        info = PARTICIPANTS_MAP.get(ep_id, {})
        app_tag = info.get("app", "TRANSCRIPT").upper().replace("-", "_")
        new_filename = f"{ep_id}_{app_tag}.docx"
        new_filepath = os.path.join(target_dir, new_filename)

        # Anonimizar en un archivo temporal o directo
        anonymize_docx_file(filepath, ep_id, new_filepath)
        
        # Eliminar archivo original si el nombre cambió
        if filename != new_filename and os.path.exists(filepath):
            try:
                os.remove(filepath)
                print(f"[{ep_id}] Anonimizado y renombrado: '{filename}' -> '{new_filename}'")
            except Exception as e:
                print(f"[{ep_id}] Anonimizado como '{new_filename}' (original conservado/bloqueado por sync: {e})")
        else:
            print(f"[{ep_id}] Anonimizado en lugar: '{new_filename}'")

def main():
    print("=== INICIANDO PROCESO DE ANONIMIZACIÓN DE ENTREVISTAS ===")
    
    # 1. Procesar data/raw/transcripts
    raw_dir = os.path.join("data", "raw", "transcripts")
    print(f"\nProcesando directorio: {raw_dir}")
    process_directory(raw_dir)

    # 2. Procesar Transcripts (raíz)
    root_transcripts = "Transcripts"
    print(f"\nProcesando directorio: {root_transcripts}")
    process_directory(root_transcripts)

    # 3. Regenerar el corpus JSON anonimizado
    print("\nRegenerando data/processed/corpus.json...")
    ingest_all(transcripts_dir=raw_dir)
    
    print("\n[OK] ANONIMIZACION COMPLETADA CON EXITO.")

if __name__ == "__main__":
    main()
