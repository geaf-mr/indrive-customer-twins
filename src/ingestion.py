"""
Module for ingesting qualitative raw research documents (.docx)
and processing them into a structured corpus.
"""

import os
import re
import json
import docx

def parse_transcript_filename(filename: str):
    """Extract participant ID and metadata from transcript filename."""
    match = re.search(r'(Ep|EP)\s*(\d+)', filename, re.IGNORECASE)
    ep_id = f"EP{int(match.group(2)):02d}" if match else "EP00"
    
    clean_name = filename.replace('.docx', '')
    
    # Identify primary app tag in filename if any
    app_tag = "UNKNOWN"
    if "INDRIVE" in filename.upper() or "INDRIVE" in filename.upper():
        app_tag = "inDrive"
    elif "YANGO" in filename.upper():
        app_tag = "Yango"
    elif "UBER" in filename.upper():
        app_tag = "Uber"

    return ep_id, clean_name, app_tag

def ingest_all(transcripts_dir: str = "data/raw/transcripts", guides_dir: str = "data/raw/guides", output_path: str = "data/processed/corpus.json"):
    """Parse all docx files into a structured JSON corpus."""
    corpus = []
    
    if not os.path.exists(transcripts_dir):
        raise FileNotFoundError(f"Directory {transcripts_dir} does not exist.")

    for filename in sorted(os.listdir(transcripts_dir)):
        if not filename.endswith(".docx"):
            continue

        filepath = os.path.join(transcripts_dir, filename)
        ep_id, clean_name, app_tag = parse_transcript_filename(filename)

        try:
            doc = docx.Document(filepath)
        except Exception as e:
            print(f"Error reading {filename}: {e}")
            continue

        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # Chunk transcript into logical dialogue snippets
        chunk_size = 3
        for i in range(0, len(paragraphs), chunk_size):
            chunk_text = "\n".join(paragraphs[i:i+chunk_size])
            if len(chunk_text) < 20:
                continue

            corpus.append({
                "id": f"{ep_id}_chunk_{i//chunk_size:03d}",
                "transcript_id": ep_id,
                "interviewee_label": clean_name,
                "app_affiliation": app_tag,
                "source_file": filename,
                "chunk_index": i // chunk_size,
                "text": chunk_text
            })

    # Process PDF Report chunks
    pdf_report_path = os.path.join(os.path.dirname(transcripts_dir), "reports", "peru_tuk_tuk_indrive_2026.json")
    if os.path.exists(pdf_report_path):
        with open(pdf_report_path, "r", encoding="utf-8") as f:
            pdf_data = json.load(f)
            for item in pdf_data:
                page_num = item.get("page", 0)
                text = item.get("text", "")
                if text:
                    corpus.append({
                        "id": f"REPORT_PDF_page_{page_num:02d}",
                        "transcript_id": f"REPORT_P{page_num:02d}",
                        "interviewee_label": f"Informe General Tuk Tuk (Página {page_num})",
                        "app_affiliation": "Informe General",
                        "source_file": "PERU TUK TUK INDRIVE.pdf",
                        "chunk_index": page_num,
                        "text": text
                    })

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(corpus, f, ensure_ascii=False, indent=2)

    print(f"Successfully processed {len(corpus)} text chunks into {output_path}")

    # Process Brief files
    brief_doc_path = os.path.join("data/raw/brief", "Brief de Mototaxis Lima.docx")
    brief_data = {}
    if os.path.exists(brief_doc_path):
        doc = docx.Document(brief_doc_path)
        brief_data["brief_text"] = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

    brief_json_path = "data/processed/brief_context.json"
    with open(brief_json_path, "w", encoding="utf-8") as f:
        json.dump(brief_data, f, ensure_ascii=False, indent=2)
    print(f"Saved brief context to {brief_json_path}")

    return corpus

if __name__ == "__main__":
    ingest_all()
